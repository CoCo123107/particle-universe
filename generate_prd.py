# -*- coding: utf-8 -*-
"""Generate PRD Word document for 3D Particle Interactive System (V2.1)."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], '1F4E79')
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('3D 粒子手势交互系统\n产品需求文档（PRD）')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('抖音 AI 创作者计划 · 情感化国风粒子互动体验')
    r.font.size = Pt(14)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run('版本：V2.1（架构评审修订版）  |  日期：2026-06-15  |  交付：HTML + 内嵌/按需模型数据')
    r.font.size = Pt(10)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()

    # ========== 1. 文档概述 ==========
    add_heading(doc, '1. 文档概述', 1)
    add_para(doc, '本文档为「3D 粒子手势交互系统」的完整产品需求说明，面向抖音 AI 创作者计划活动场景。V2.1 在 V2.0 基础上，依据 WebGL/Three.js 架构评审结论，修订了 Morph 算法、粒子动力学、特效状态机、性能分级、资源生命周期与 UI 状态优先级，确保方案可落地、可验收、移动端可运行。')
    add_para(doc, '交付形态：主文件 particle-universe.html（可独立运行）+ models/ 压缩点集（内嵌或按需 fetch），Three.js / MediaPipe 走 CDN。')

    add_heading(doc, '1.1 文档目的', 2)
    add_bullet(doc, '明确可执行的功能边界、交互逻辑与技术实现方案（无逻辑自相矛盾）')
    add_bullet(doc, '为 Cursor / Trae 等 AI 工具提供分模块、可验证的生成规格')
    add_bullet(doc, '定义分档性能指标与资源 dispose 验收标准')
    add_bullet(doc, '支撑抖音创作者计划的「可玩性 + 共鸣 + 文化传播」传播诉求')

    add_heading(doc, '1.2 目标用户', 2)
    add_table(doc, ['用户类型', '核心诉求', '使用场景'], [
        ['抖音创作者', '录制高互动、高完播短视频', '竖屏 9:16 面对摄像头互动'],
        ['普通玩家', '轻松上手、视觉震撼、有情感反馈', '浏览器打开即玩'],
        ['文化爱好者', '感受国风元素与节日氛围', '切换敦煌、青花瓷、生肖等模型'],
        ['活动评审/观众', '30 秒内看到核心 Hook', '双手张合 → 粒子响应'],
    ])

    add_heading(doc, '1.3 产品定位', 2)
    add_para(doc, '一句话定位：用「双手 + 情感 + 国风」驱动的实时 3D 粒子宇宙，让每一次张合都是心跳，每一次手势都是魔法。')
    add_bullet(doc, '技术标签：Three.js + MediaPipe HandLandmarker + 单 Scene 单 Points 架构')
    add_bullet(doc, '体验标签：低延迟手势、极角排序 Morph、半透明发光星云')
    add_bullet(doc, '传播标签：国风文创、节日限定、文字粒子化、手势彩蛋')

    add_heading(doc, '1.4 术语表', 2)
    add_table(doc, ['术语', '定义'], [
        ['统一粒子数 N', '全局固定 6000（可降级至 2500/4000，见 §5.5）'],
        ['baseTarget[i]', '粒子 i 在当前模型上的归一化目标点（相对质心）'],
        ['极角排序', '离线按 atan2(y,x) 对点集排序，保证跨模型 morph 轨迹连续'],
        ['Cross-fade Morph', '淡出噪声 → 位置插值 → 淡入的三段式切换'],
        ['spread', '双手距离归一化标量 0~1，0=合拢，1=张开'],
        ['effectOffset', '特效产生的位移偏移，Verlet 积分 + 弹簧回弹'],
        ['userColor', '用户 color picker 选定色，特效结束后 200ms lerp 恢复'],
        ['chargeState', '蓄力 FSM：idle | charging | armed | bursting'],
    ])

    add_heading(doc, '1.5 V2.1 修订摘要（相对 V2.0）', 2)
    add_table(doc, ['原问题', '修订方案', '涉及章节'], [
        ['索引 morph 视觉混乱', '极角排序 + cross-fade，禁止匈牙利算法', '§4.3、§7.3'],
        ['velocity burst 无 shader 支撑', '统一 Verlet 动力学 + effectOffset', '§4.1、§5.4'],
        ['OK 手势判定写反', '修正为「三指弯曲」', '§8.3'],
        ['spread 依赖 face scale', '前 60 帧 auto-calibrate min/max', '§8.5'],
        ['chargeState 与 FSM 不一致', '统一四态 + activeEffect 从属', '§7.1、§7.2'],
        ['800KB 与 18 模型点集矛盾', '模型 gzip 分文件按需加载', '§2.2、§5.6'],
        ['无 dispose 规范', '新增 §5.7 生命周期 Checklist', '§5.7'],
        ['颜色/节日/特效覆盖冲突', 'userColor 为主 + effectTint 乘法 + lerp 回退', '§7.4'],
        ['全屏 resize 未定义', '新增 §7.5 Resize 规范', '§7.5'],
        ['移动端 280px 侧栏', '≤768px 改为底部抽屉', '§6.1'],
        ['默认 bloom+curl 掉帧', '三档性能，默认 Mid 无 bloom', '§5.5'],
    ])

    # ========== 2. 产品目标 ==========
    add_heading(doc, '2. 产品目标与成功指标', 1)

    add_heading(doc, '2.1 业务目标', 2)
    add_bullet(doc, '满足抖音 AI 创作者计划对「可交互、可共鸣、文化传播」的要求')
    add_bullet(doc, '30 秒内可展示 Hook：摄像头 + 粒子随双手张合')
    add_bullet(doc, '嵌入国风文创元素，支持文字粒子化 UGC')
    add_bullet(doc, '移动端 WebView 可玩（性能档 Low 仍 ≥30fps）')

    add_heading(doc, '2.2 体验与性能指标', 2)
    add_table(doc, ['指标', '目标值', '测量方式'], [
        ['手势到视觉反馈', '≤ 150ms（Mid 档）', 'landmarks 更新 → uniform 生效 → render 帧'],
        ['帧率 Mid 档', '≥ 45 FPS（1080p 桌面中端）', 'rAF 计数 10s 平均'],
        ['帧率 Low 档', '≥ 30 FPS（手机竖屏 WebView）', '同上，抖音内置浏览器'],
        ['Morph 过渡', '1.5s cross-fade，无穿模交叉', '肉眼轨迹连续、无大面积横穿'],
        ['首次可交互', '≤ 5s（含 WASM + 首模型）', '授权后 spread 可控制'],
        ['主 HTML 体积', '≤ 200KB gzip（不含模型与 CDN 库）', '模型独立按需加载'],
        ['10 分钟内存', '堆增长 < 5MB', 'Performance.memory 或 DevTools'],
    ])

    add_heading(doc, '2.3 非目标（Out of Scope）', 2)
    add_bullet(doc, '不依赖后端、数据库、用户账号')
    add_bullet(doc, '不做多人联网、AR 眼镜专用适配')
    add_bullet(doc, '不做离线匈牙利匹配、不做 per-particle curl noise（Mid/Low 档）')
    add_bullet(doc, '不内置 3D 建模工具；轮廓来自 SVG 离线采样')

    # ========== 3. 用户故事 ==========
    add_heading(doc, '3. 用户故事与使用流程', 1)

    add_heading(doc, '3.1 核心用户故事', 2)
    add_table(doc, ['角色', '需求', '价值'], [
        ['创作者', '双手张合控制粒子扩散', '录制「手控宇宙」短视频'],
        ['玩家', '切换模型时粒子平滑变形', '视觉连贯'],
        ['玩家', '双手靠近心跳加速', '情感共鸣'],
        ['玩家', '快合快开触发宇宙爆炸', '惊喜成就感'],
        ['玩家', 'OK/点赞/比心触发魔法', '探索隐藏玩法'],
        ['文化爱好者', '敦煌、青花瓷等国风轮廓', '文化传播'],
        ['玩家', '输入汉字粒子化', '个性化祝福'],
        ['玩家', '调色实时预览且特效后不丢色', '主题一致'],
    ])

    add_heading(doc, '3.2 主流程（Happy Path）', 2)
    for i, step in enumerate([
        '打开页面 → 星空 + 默认星云粒子',
        '授权摄像头 → 前 60 帧自动校准 spread → 显示「手势已连接」',
        '双手张合 → 粒子随 spread 缩放/扩散',
        '选择「敦煌飞天」→ 1.5s cross-fade morph',
        '输入「福」→ 300ms debounce 后 morph 为文字（模型卡片仍高亮，文字优先显示）',
        '清空文字 → 回到上次选中模型',
        '快合快开 → 宇宙爆炸 → 2.5s 后 spring 回 baseTarget + userColor',
        '比心 → 爱心特效 3s，不阻塞蓄力（见 §7.2）',
        '全屏 → canvas resize + 侧栏/抽屉折叠',
    ], 1):
        add_bullet(doc, f'{i}. {step}')

    add_heading(doc, '3.3 异常流程', 2)
    add_table(doc, ['场景', '系统行为', '用户提示'], [
        ['拒绝摄像头', '鼠标 Y 拖拽 / 触控降级', 'Toast：已切换手动模式'],
        ['未检测到手', 'spread 保持最后值，noise 流动继续', '「请将双手放入画面」'],
        ['低 FPS', '自动降档 Mid→Low', '设置显示当前性能档'],
        ['Morph 中再切换', '从当前 render 位置重新 cross-fade', '无阻塞'],
        ['字体未加载', '文字 morph 排队', '输入框旁 spinner'],
        ['WebGL 上下文丢失', '重建 renderer，复用 AppState', 'Toast：正在恢复…'],
        ['visibility 隐藏', 'pause rAF + stop camera tracks', '回前台自动 resume'],
    ])

    # ========== 4. 功能需求 ==========
    add_heading(doc, '4. 功能需求详述', 1)

    add_heading(doc, '4.1 基础粒子系统（P0）', 2)

    add_para(doc, '4.1.1 视觉规格', bold=True)
    add_bullet(doc, '统一粒子数 N：6000（High），4000（Mid，默认），2500（Low）')
    add_bullet(doc, '外观：半透明发光 Points，Additive Blending，soft circle sprite')
    add_bullet(doc, '背景：单层/双层静态星空（Mid/Low 无 parallax）；High 可选慢速 parallax')
    add_bullet(doc, '流动：Mid/Low 用 sin 相位偏移；High 可用轻量 2D noise（禁止 curl noise）')
    add_bullet(doc, '深度：size attenuation；Mid/Low 不用 Scene fog（省 fill-rate）')

    add_para(doc, '4.1.2 粒子动力学（CPU 主循环，每帧 O(N)）', bold=True)
    add_para(doc,
        '统一状态方程（实现必须按此顺序）：\n'
        '1) base = centroid + baseTarget[i] * spreadScale(spread) * heartbeatScale(bpm, t)\n'
        '2) noise = simpleNoise(i, t) * noiseAmp * (1 - morphBlend)  // morph 期间 noise 淡出\n'
        '3) pos = base + noise + effectOffset[i]\n'
        '4) effectOffset：Verlet 积分；爆炸时 radialImpulse×8；每帧 springTo(base) + damping(0.92)\n'
        '5) 写回 position attribute；禁止在 fragment 阶段改位置\n'
        'spreadScale(s) = 1 + s * 0.5；heartbeatScale = 1 + 0.06*sin(2π*BPM/60*t)（仅合拢时显著）\n'
        'centroid：当前模型点集质心，保证缩放绕形心而非世界原点')

    add_para(doc, '4.1.3 双手张合控制（P0）', bold=True)
    add_bullet(doc, '输入：MediaPipe HandLandmarker，最多 2 手，21 关键点')
    add_bullet(doc, '度量：双手腕 landmark[0] 欧氏距离（归一化图像坐标）')
    add_bullet(doc, '校准：启动后前 60 帧记录 distMin/distMax，spread = clamp((d-distMin)/(distMax-distMin), 0, 1)')
    add_bullet(doc, '滤波：One Euro Filter（mincutoff=1.0, beta=0.007）')
    add_bullet(doc, '无手时：spread 冻结，不重置校准值')

    add_para(doc, '4.1.4 模型 Morph（P0）', bold=True)
    add_bullet(doc, 'Cross-fade 三阶段（总 1.5s）：')
    add_bullet(doc, '  · 0~0.3s：morphBlend 1→0.7，noiseAmp 线性降至 0', level=1)
    add_bullet(doc, '  · 0.3~1.2s：baseTarget 从 A 极角序 lerp 到 B 极角序，morphT 0→1', level=1)
    add_bullet(doc, '  · 1.2~1.5s：morphBlend 0.7→1，noiseAmp 恢复', level=1)
    add_bullet(doc, '中断切换：snapshot 当前 pos 作为新 A，立即重新 cross-fade')
    add_bullet(doc, 'Morph 期间 spread/heartbeat 仍生效（作用于 base 项）')

    add_para(doc, '4.1.5 颜色选择器（P0）', bold=True)
    add_bullet(doc, '单色 <input type="color"> + HEX；oninput 更新 userColor uniform')
    add_bullet(doc, '预设：故宫红、天青、琉璃黄、墨黑、星蓝')
    add_bullet(doc, '最终色：finalColor = userColor * effectTint（见 §7.4）')

    add_para(doc, '4.1.6 UI 与全屏（P0）', bold=True)
    add_bullet(doc, '桌面：右侧固定面板 280px；canvas = calc(100vw - 280px)')
    add_bullet(doc, '移动 ≤768px：底部抽屉 45vh，canvas 占满宽；FAB 展开面板')
    add_bullet(doc, '全屏：Fullscreen API；见 §7.5 resize')
    add_bullet(doc, '扁平化、8px 圆角；Mid/Low 档面板不用 backdrop-filter（省 GPU）')

    add_heading(doc, '4.2 情感化交互系统（P0）', 2)

    add_para(doc, '4.2.1 心跳模式', bold=True)
    add_bullet(doc, '默认开启；面板可关')
    add_bullet(doc, 'BPM = lerp(72, 132, 1 - spread)  // 合拢越快，上限略降避免晕眩')
    add_bullet(doc, 'heartbeatScale 仅当 spread < 0.35 时全幅生效；张开时减弱至 0.02')
    add_bullet(doc, '面板显示 BPM + 心跳图标')

    add_para(doc, '4.2.2 能量蓄力 · 宇宙爆炸', bold=True)
    add_bullet(doc, '检测：400ms 内 spread 从 >0.65 降至 <0.25，再 600ms 内升至 >0.75')
    add_bullet(doc, 'charging：effectTint 偏暖 (1.2, 1.0, 0.85)，粒子 effectOffset 向 centroid 聚拢')
    add_bullet(doc, 'armed：UI 蓄力条满，等待张开')
    add_bullet(doc, 'bursting：0~300ms radialImpulse×8；300~1200ms 全屏 shockwave（单 Mesh uniform，不 new Points）')
    add_bullet(doc, '1200~2500ms：spring + damping 回 base；effectTint lerp 回 (1,1,1)')
    add_bullet(doc, '冷却 3s；bursting 期间忽略新的蓄力')

    add_para(doc, '4.2.3 手势魔法', bold=True)
    add_table(doc, ['手势', '识别条件', '特效', '实现约束', '时长'], [
        ['OK', '拇食捏合，中指/无名/小指弯曲', '时空涟漪', 'shockwave uniform 同心纹', '2s'],
        ['点赞', '拇指伸展，四指弯曲，单手', '星光祝福', 'effectOffset Y+ 偏置 + 金色 tint', '2.5s'],
        ['比心', '双手拇食形成心形', '爱心降临', 'tint 粉色 + 临时 heart overlay 点集 lerp', '3s'],
        ['剪刀手', '食+中伸展（P1）', '双子分裂', 'centroid 临时分裂为双质心', '2s'],
    ])
    add_bullet(doc, '几何规则识别，无需 ML 分类器')
    add_bullet(doc, '防抖 2s；详见 §7.2 优先级')

    add_heading(doc, '4.3 国风文创元素库（P0）', 2)

    add_para(doc, '4.3.1 模型清单（均离线预处理为 N=6000 极角序点集）', bold=True)
    add_table(doc, ['分类', 'ID', '显示名'], [
        ['国风', 'dunhuang_feitian', '敦煌飞天'],
        ['国风', 'qinghua_ci', '青花瓷'],
        ['国风', 'shengxiao_dragon', '生肖·龙'],
        ['国风', 'peking_mask', '京剧脸谱'],
        ['国风', 'gu_shan', '古风扇面'],
        ['国风物件', 'dragon', '龙'],
        ['国风物件', 'panda', '熊猫'],
        ['国风物件', 'tuan_shan', '团扇'],
        ['国风物件', 'lantern', '灯笼'],
        ['经典', 'heart', '爱心'],
        ['经典', 'saturn', '土星'],
        ['经典', 'rose', '玫瑰'],
        ['经典', 'firework', '烟花'],
        ['经典', 'gift_box', '礼物盒'],
        ['节日', 'spring_firework', '春节烟花'],
        ['节日', 'mid_autumn_rabbit', '中秋玉兔'],
        ['节日', 'qixi_bridge', '七夕鹊桥'],
        ['动态', 'hanzi', '汉字（用户输入）'],
    ])

    add_para(doc, '4.3.2 点集生成（离线）', bold=True)
    add_bullet(doc, 'SVG 512×512 → 均匀采样 6000 点 → 去质心 → 归一化 max(r)=1')
    add_bullet(doc, '按 atan2(y,x) 升序排序，写入 models/{id}.json.gz')
    add_bullet(doc, '汉字：opentype.js 离线或运行时 getPath → 每字 ceil(6000/字数) 点 → 合并后 resample 至 6000')
    add_bullet(doc, '运行时降级粒子数：对 baseTarget 等距 subsample 同一排序索引')

    add_para(doc, '4.3.3 文字输入（P0）', bold=True)
    add_bullet(doc, '1~8 字符；debounce 300ms')
    add_bullet(doc, '优先级：customText 非空 → 显示文字点集；清空 → 恢复 selectedModelId')
    add_bullet(doc, '选模型不清空文字；需用户点「清除文字」或删空输入框')
    add_bullet(doc, 'fontStatus: loading | ready | error；loading 时排队 morph')

    add_para(doc, '4.3.4 节日限定', bold=True)
    add_bullet(doc, '农历节日自动高亮对应模型卡片（可手动切）')
    add_bullet(doc, '节日仅推荐 presetColor，不强制覆盖 userColor（见 §7.4）')
    add_bullet(doc, '中秋：背景可选静态 moon glow 贴图（Low 档跳过）')

    add_heading(doc, '4.4 辅助功能（P1）', 2)
    add_bullet(doc, '手动模式：鼠标在 canvas 上 drag Y → spread；touch 单指 Y + 双指 pinch')
    add_bullet(doc, 'pointer capture 绑定 canvas，避免与面板 scroll 冲突')
    add_bullet(doc, 'FPS / 延迟 / 性能档开发者 toggle')
    add_bullet(doc, 'canvas.toDataURL 截图')

    # ========== 5. 技术架构 ==========
    add_heading(doc, '5. 技术架构与实现方案', 1)

    add_heading(doc, '5.1 技术栈', 2)
    add_table(doc, ['层级', '选型', '用途'], [
        ['渲染', 'Three.js r160+', '单 Scene、单 Points、可选 EffectComposer'],
        ['手势', '@mediapipe/tasks-vision HandLandmarker', '双手关键点'],
        ['字体', 'opentype.js + Noto Serif SC', '汉字轮廓'],
        ['UI', '原生 HTML/CSS', '无框架'],
        ['模型', 'fetch models/*.json.gz', '按需加载 + Cache API'],
    ])

    add_heading(doc, '5.2 架构原则', 2)
    add_bullet(doc, '单 Scene、单 Points、单 BufferGeometry、单 PointsMaterial（全生命周期）')
    add_bullet(doc, '特效只改 uniform 与 effectOffset CPU 数组，禁止 trigger 时 new Points/Mesh（shockwave 除外单 Mesh）')
    add_bullet(doc, 'Morph 只更新 baseTarget attribute，不 recreate geometry')
    add_bullet(doc, 'AppState 为唯一真相源；UI 只写 AppState，RenderLoop 只读')

    add_heading(doc, '5.3 模块划分', 2)
    add_table(doc, ['模块', '职责', '接口'], [
        ['HandTracker', '摄像头 + HandLandmarker', 'start(), getLandmarks(), pause(), resume()'],
        ['SpreadCalibrator', '前 60 帧 min/max', 'update(d) → spread'],
        ['GestureEngine', 'spread/BPM/蓄力/手势', 'update(lm) → GestureState'],
        ['ModelRegistry', 'gzip 点集 + 汉字', 'load(id), getTargets(id|text)'],
        ['MorphController', 'cross-fade + snapshot', 'transitionTo(targets), cancel()'],
        ['ParticleSystem', 'Verlet + attribute 写回', 'tick(dt, state)'],
        ['EffectOverlay', 'shockwave 单 Mesh', 'triggerBurst(), update(dt), dispose()'],
        ['UIManager', '面板/全屏/颜色', 'bind(store), onResize()'],
        ['Lifecycle', 'dispose/pause', 'destroy(), onVisibilityChange()'],
    ])

    add_heading(doc, '5.4 Shader 规格（简化，位置在 CPU）', 2)
    add_para(doc, 'Vertex Shader：')
    add_bullet(doc, 'attribute vec3 position; // CPU 已算好最终 pos')
    add_bullet(doc, 'attribute float aMorphBlend; // 可选，控制 sprite 大小脉动')
    add_bullet(doc, 'uniform vec3 uColor; uniform float uOpacity, uPulse;')
    add_bullet(doc, 'gl_PointSize = uBaseSize * (280.0 / -mvPosition.z);')
    add_para(doc, 'Fragment Shader：')
    add_bullet(doc, 'soft circle alpha；finalColor = uColor * uPulse')
    add_bullet(doc, 'AdditiveBlending；Mid/Low 关闭 post-processing Bloom')

    add_heading(doc, '5.5 性能分级（验收必测）', 2)
    add_table(doc, ['档位', '粒子 N', 'Bloom', 'Noise', '默认场景'], [
        ['High', '6000', '可选', '2D noise', '桌面用户手动选'],
        ['Mid', '4000', '关', 'sin 相位', '桌面默认'],
        ['Low', '2500', '关', '关', '移动/WebView 默认，FPS<30 自动降入'],
    ])
    add_bullet(doc, '降级时：对 baseTarget 等距 subsample，索引 0..N-1 连续')
    add_bullet(doc, 'MediaPipe：modelComplexity=1；detectionConfidence=0.6')
    add_bullet(doc, 'setPixelRatio(min(dpr, 1.5)) on Low')

    add_heading(doc, '5.6 交付文件结构', 2)
    add_para(doc,
        'particle-universe.html          # 主逻辑 ≤200KB gzip\n'
        'models/\n'
        '  heart.json.gz                 # 每文件 ~30-50KB gzip\n'
        '  dunhuang_feitian.json.gz\n'
        '  ...                           # 按需 fetch，Cache API 缓存\n'
        'assets/noto-serif-sc-subset.woff2  # 常用字子集，可选')

    add_heading(doc, '5.7 资源生命周期 Checklist（强制）', 2)
    add_bullet(doc, '初始化：创建一次 geometry/material/points，add 到 scene')
    add_bullet(doc, '模型切换：仅 geometry.attributes.position.array.set + needsUpdate')
    add_bullet(doc, 'EffectOverlay Mesh：burst 结束 hidden=true，不 dispose 直到 destroy')
    add_bullet(doc, 'destroy()：cancelAnimationFrame；video.srcObject.getTracks().forEach stop；geometry.dispose；material.dispose；renderer.dispose()')
    add_bullet(doc, 'visibilitychange hidden：pause rAF + HandTracker.pause()')
    add_bullet(doc, 'visibilitychange visible：resume；若 webgl context lost 则重建 renderer 并 re-upload attributes')
    add_bullet(doc, '禁止：每次 morph new BufferGeometry；禁止 burst 每帧 new Material')

    # ========== 6. UI ==========
    add_heading(doc, '6. UI/UX 设计规范', 1)

    add_heading(doc, '6.1 布局', 2)
    add_table(doc, ['断点', '布局', 'canvas'], [
        ['≥769px', '右栏 280px fixed', 'calc(100vw - 280px) × 100vh'],
        ['≤768px', '底部抽屉 + FAB', '100vw × calc(100vh - 56px)'],
        ['全屏', '面板折叠为角标', '100vw × 100vh'],
    ])

    add_heading(doc, '6.2 视觉', 2)
    add_table(doc, ['元素', '规范'], [
        ['背景', '#0a0e17'],
        ['面板', 'rgba(15,20,35,0.92)；Mid/Low 无 blur'],
        ['accent', '#6c9eff；国风 #c41e3a'],
        ['按钮', '36px 高，8px 圆角'],
    ])

    add_heading(doc, '6.3 状态反馈', 2)
    add_bullet(doc, '模型卡片 active 态；文字模式时模型仍 active 但顶栏 Tag「文字预览」')
    add_bullet(doc, '蓄力 armed：边缘 vignette + 蓄力条')
    add_bullet(doc, '手势触发：emoji popup 1.5s')

    # ========== 7. 状态机 ==========
    add_heading(doc, '7. 数据结构与状态机', 1)

    add_heading(doc, '7.1 AppState（唯一真相源）', 2)
    add_para(doc,
        '{\n'
        '  selectedModelId: "heart",\n'
        '  customText: "",\n'
        '  displayMode: "model" | "text",  // text 仅当 customText 非空\n'
        '  userColor: "#6c9eff",\n'
        '  effectTint: [1,1,1],\n'
        '  spread: 0.5,\n'
        '  bpm: 72,\n'
        '  chargeState: "idle"|"charging"|"armed"|"bursting",\n'
        '  gestureEffect: null|"ok"|"thumbsup"|"heart"|"scissors",\n'
        '  morph: { phase:"idle"|"fadeOut"|"lerp"|"fadeIn", t, fromTargets, toTargets, snapshot[] },\n'
        '  handDetected: false,\n'
        '  perfTier: "mid"|"low"|"high",\n'
        '  fontStatus: "idle"|"loading"|"ready"|"error",\n'
        '  spreadCalib: { min, max, frames, done }\n'
        '}')

    add_heading(doc, '7.2 蓄力与手势优先级 FSM', 2)
    add_table(doc, ['优先级', '状态', '规则'], [
        ['1', 'bursting', '独占；忽略蓄力重入；gestureEffect 可排队至 bursting 结束'],
        ['2', 'gestureEffect 激活', '可打断 charging/armed，chargeState 重置 idle'],
        ['3', 'charging → armed', '正常蓄力'],
        ['4', 'idle + spread', '基础张合与心跳'],
    ])
    add_bullet(doc, 'gestureEffect 与 bursting 互斥：bursting 期间手势只显示 UI 提示「爆炸中」')
    add_bullet(doc, '同一手势 2s 防抖')

    add_heading(doc, '7.3 Morph 算法（定稿）', 2)
    add_bullet(doc, '离线：所有模型 resample + 极角排序 → 6000×3 Float32')
    add_bullet(doc, '运行时 transitionTo(B)：')
    add_bullet(doc, '  snapshot = 当前 position array 拷贝', level=1)
    add_bullet(doc, '  phase fadeOut 0.3s：noiseAmp→0', level=1)
    add_bullet(doc, '  phase lerp 0.9s：lerp(snapshot[i], B[i], easeInOutCubic(t))', level=1)
    add_bullet(doc, '  phase fadeIn 0.3s：noiseAmp 恢复', level=1)
    add_bullet(doc, '完成：baseTarget=B，snapshot 释放')

    add_heading(doc, '7.4 颜色与特效优先级', 2)
    add_bullet(doc, 'userColor：picker 唯一写入源，持久保存')
    add_bullet(doc, 'effectTint：特效临时乘色；charging (1.2,1.0,0.85)；比心 (1.2,0.85,0.95)；点赞 (1.15,1.05,0.8)')
    add_bullet(doc, 'finalColor = userColor * effectTint；特效结束 200ms lerp effectTint→(1,1,1)')
    add_bullet(doc, '节日 preset：仅点击「应用节日色」写入 userColor，不自动覆盖')
    add_bullet(doc, '禁止：特效 hardcode 替换 userColor 且不恢复')

    add_heading(doc, '7.5 Resize 与全屏', 2)
    add_bullet(doc, 'window.resize + fullscreenchange → Layout.compute() → renderer.setSize(w,h) → camera.aspect=w/h → updateProjectionMatrix')
    add_bullet(doc, '移动端 address bar 变化：监听 visualViewport.resize')
    add_bullet(doc, '全屏进入：panel 折叠；spread 手势不受影响')
    add_bullet(doc, '手动模式 pointer：坐标 relative to canvas.getBoundingClientRect()')

    # ========== 8. 手势 ==========
    add_heading(doc, '8. 手势识别', 1)

    add_heading(doc, '8.1 关键点', 2)
    add_para(doc, '0:WRIST, 4:THUMB_TIP, 8:INDEX_TIP, 12:MIDDLE_TIP, 16:RING_TIP, 20:PINKY_TIP')

    add_heading(doc, '8.2 辅助函数', 2)
    add_bullet(doc, 'isFingerExtended(tip,pip,wrist)：tip 距 wrist > pip 距 wrist × 1.1')
    add_bullet(doc, 'dist(a,b)：2D 欧氏距离（归一化坐标）')

    add_heading(doc, '8.3 判定规则（V2.1 修正）', 2)
    add_para(doc,
        'OK: dist(thumbTip,indexTip) < 0.05 && !extended(middle) && !extended(ring) && !extended(pinky)\n'
        'ThumbsUp: extended(thumb) && !extended(index,middle,ring,pinky)\n'
        'Heart: 双手 indexTip 相距适中，两 thumbTip 靠近，形成心形包络\n'
        'Spread: SpreadCalibrator.update(dist(wristL,wristR))')

    add_heading(doc, '8.4 滤波', 2)
    add_bullet(doc, 'One Euro on spread；BPM 由 spread 导出，不再单独滤波')

    add_heading(doc, '8.5 Spread 自动校准', 2)
    add_bullet(doc, '启动后 60 帧或检测到 dist 稳定变化后 done=true')
    add_bullet(doc, 'distMin = min observed * 0.95；distMax = max observed * 1.05')
    add_bullet(doc, 'done 前 UI 显示「正在校准手势…」')

    # ========== 9. 模型资源 ==========
    add_heading(doc, '9. 模型资源规范', 1)
    add_bullet(doc, 'JSON 格式：{ id, name, category, n:6000, points:Float32Array base64 或数组 }')
    add_bullet(doc, 'gzip 后单模型 30~50KB；首屏仅加载 heart.json.gz + 当前选中')
    add_bullet(doc, 'build 脚本：sample-svg.js → sort-polar.js → gzip → models/')

    # ========== 10. 开发计划 ==========
    add_heading(doc, '10. 开发计划', 1)
    add_table(doc, ['阶段', '周期', '交付', '验收'], [
        ['M1', '2d', '单 Points + spread 校准 + 手动模式', 'Mid 档 45fps'],
        ['M2', '2d', '极角 morph + 5 模型按需加载', '切换无横穿'],
        ['M3', '2d', '蓄力 FSM + 3 手势 + 颜色 lerp', '特效后色恢复'],
        ['M4', '2d', '全部国风 + 文字 + 生命周期', '10min 无 leak'],
        ['M5', '1d', '移动抽屉 + 三档性能 + 全屏', 'Low 档 30fps'],
    ])

    add_heading(doc, '10.1 Cursor 生成 Prompt（V2.1）', 2)
    add_para(doc,
        '「按 PRD V2.1 实现 particle-universe.html。硬性约束：单 Scene 单 Points；'
        'CPU Verlet 动力学；极角排序 cross-fade morph；spread 60 帧校准；'
        'chargeState 四态 FSM；userColor*effectTint；Mid 档 4000 粒子无 bloom；'
        '模型 fetch models/*.json.gz；destroy/visibility 生命周期；'
        '桌面右栏/移动底抽屉；OK 手势三指弯曲。禁止 curl noise、禁止匈牙利算法、'
        '禁止 morph 时 new Geometry。」')

    # ========== 11. 测试 ==========
    add_heading(doc, '11. 测试验收', 1)

    add_table(doc, ['ID', '用例', '预期'], [
        ['TC-01', '双手张合', 'spread 与幅度成正比，校准后满量程'],
        ['TC-02', 'heart→dragon morph', '1.5s 连续，无大面积横穿'],
        ['TC-03', '调色后爆炸', '结束后 200ms 回到 userColor'],
        ['TC-04', '输入福字再清空', '回 selectedModelId 轮廓'],
        ['TC-05', 'OK 手势', '三指弯曲捏合可触发'],
        ['TC-06', 'burst 冷却', '3s 内不重复'],
        ['TC-07', '全屏 resize', '粒子不变形，spread 正常'],
        ['TC-08', '10min 交互', '内存增长 <5MB'],
        ['TC-09', 'Low 档手机', '≥30fps'],
        ['TC-10', 'visibility 切换', '无 leak，恢复可玩'],
    ])

    add_heading(doc, '11.2 兼容性', 2)
    add_table(doc, ['环境', '要求'], [
        ['Chrome/Edge 90+', '完整功能，推荐'],
        ['Safari 15+', 'HTTPS；本地用 live server'],
        ['抖音 WebView', '默认 Low 档，竖屏'],
        ['file://', '摄像头可能不可用；提示部署到 HTTPS'],
    ])

    # ========== 12. 风险 ==========
    add_heading(doc, '12. 风险与应对', 1)
    add_table(doc, ['风险', '应对'], [
        ['Morph 仍不够像', '增加 cross-fade 至 2s；heart→复杂模型加中间态 nebula'],
        ['MediaPipe 延迟', 'render 用 spread 预测外推 1 帧'],
        ['字体加载失败', 'fallback 系统 sans；显示 error 态'],
        ['WebGL context lost', '重建 renderer + re-upload attributes'],
        ['模型 fetch 失败', 'retry 2 次 + 回退 heart'],
    ])

    add_heading(doc, '13. 版本记录', 1)
    add_table(doc, ['版本', '日期', '变更'], [
        ['V1.0', '2026-06-01', '初版'],
        ['V2.0', '2026-06-15', '情感交互 + 国风库'],
        ['V2.1', '2026-06-15', '架构评审修订：Morph/FSM/性能/生命周期/颜色优先级'],
    ])

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('— 文档结束 —')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(150, 150, 150)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return doc

if __name__ == '__main__':
    out_dir = os.path.dirname(os.path.abspath(__file__))
    for name in ['PRD_3D粒子手势交互系统_V2.1.docx', 'PRD_3D_Particle_Interactive_System_V2.1.docx']:
        path = os.path.join(out_dir, name)
        build_document().save(path)
        print(f'Generated: {path}')
